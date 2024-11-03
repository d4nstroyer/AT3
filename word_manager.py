# -*- coding: utf-8 -*-
"""
Created on Wed Oct 30 10:52:56 2024

@author: Danny
"""

import os
from docx import Document
from docx.shared import Inches

# Function to create a Word document
def create_file(filename):
    doc = Document()
    
    # Create two headings and their subheadings in the expected order
    doc.add_heading('Heading 1', level=1)
    doc.add_heading('Subheading under Heading 1', level=2)
    doc.add_paragraph('This is a paragraph under Heading 1.')
    
    doc.add_heading('Heading 2', level=1)
    doc.add_heading('Subheading under Heading 2', level=2)
    doc.add_paragraph('This is a paragraph under Heading 2.')

    # Add a picture to the Word document 
    if os.path.exists('image.jpg'):  
        doc.add_picture('image.jpg', width=Inches(2))
    else:
        print("Image file upload error.")

    # Save the document
    doc.save(filename)
    print(f"Document '{filename}' created successfully.")

# Function to get full text from a Word document
def get_full_text(filename):
    if not os.path.exists(filename):
        print(f"File '{filename}' does not exist.")
        return ""

    doc = Document(filename)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def main():
    while True:
        print("\nWord Document Manager - Choose an option:")
        print("1. Create a Word document")
        print("2. Get full text from a Word document")
        print("3. Exit")
        
        choice = input("Enter your choice (1-3): ")

        if choice == '1':
            filename = "document.docx"
            create_file(filename)
        
        elif choice == '2':
            filename = input("Enter the filename to read (with .docx extension): ")
            full_text = get_full_text(filename)
            print("\nFull text from the document:")
            print(full_text)

        elif choice == '3':
            print("Exiting the program. Goodbye!")
            break
        
        else:
            print("Invalid choice. Please enter a valid option.")

if __name__ == "__main__":
    main()
